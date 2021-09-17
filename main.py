import os
from pathlib import Path
import docx
import re

def eachPath(filepath):
    #显示文件文件夹里所有的文件类型并且做格式判断
    pathDir = os.listdir(filepath)
    #print('pathDir:', pathDir)
    for allDir in pathDir:
        #获取文件名称拼接得到路径
        child = os.path.join(filepath, allDir)
        #Path(filepath).suffix得到后缀
        formatJudger(child,Path(child).suffix)

def formatJudger(filepath, suffix):
    #后缀列表，只能处理列出格式的解析问题
    formatList=['.pdf','.doc','.docx','.wps']
    if suffix in formatList:
        if suffix == '.pdf':
            pdfParser(filepath)
        elif suffix == '.doc':
            docParser(filepath)
        elif suffix == '.docx':
            docxParser(filepath)
        elif suffix == '.wps':
            wpsParser(filepath)
    else:
        #无法处理打印错误信息
        format=suffix.replace(".","")
        print("无法解析格式为"+format+"的文件")

def docConvert():
    ''' 一个例子
    import sys
    import pickle
    import re
    import  codecs
    import string
    import shutil
    from win32com import client as wc
    import docx

    word = wc.Dispatch('Word.Application')
    doc = word.Documents.Open(u'E:\code\\xxxx.doc')        # 目标路径下的文件
    doc.SaveAs(u'E:\\code\\hhhhhhhh.docx', 12, False, "", True, "", False, False, False, False)  # 转化后路径下的文件
    doc.Close()
    word.Quit()

    '''

def pdfParser(filepath):
    print(filepath)

def docParser(filepath):
    #doc文件无法直接处理，要转为docx,macOs无法实现，需要在windows上实现转换
    print(filepath)



def docxParser(filepath):
    print(filepath)
    doc = docx.Document(filepath)
    print(len(doc.paragraphs))

    #标题检测
    for p in doc.paragraphs:
        docxParaParser(p)


def docxParaParser(para):

    print("段落正文",para.text,sep=":")
    # 段落格式
    # 缩进,正为首行缩进，负为悬挂缩进
    if para.paragraph_format.first_line_indent is None:
        print(para.paragraph_format.first_line_indent)
    else:
        if(para.paragraph_format.first_line_indent.pt)>0:
            print("首行缩进",para.paragraph_format.first_line_indent.pt,sep=":")
        else:
            print("悬挂缩进",para.paragraph_format.first_line_indent.pt,sep=":")
    #

    # 段落对齐方式
    if para.paragraph_format.alignment is None:
        print("未设置居中格式")
    else:
        print(para.paragraph_format.alignment)

    # 段前间距
    if para.paragraph_format.space_before is None:
        print("段前间距",para.paragraph_format.space_before,sep=":")
    else:
        print("段前间距",para.paragraph_format.space_before.pt,sep=":")

    #段后间距
    if para.paragraph_format.space_after is None:
        print("段后间距", para.paragraph_format.space_after,sep=":")
    else:
        print("段后间距", para.paragraph_format.space_after.pt,sep=":")

    #行间距
    if para.paragraph_format.line_spacing is None:
        print("行间距",para.paragraph_format.line_spacing,sep=":")
    else:
        print("行间距",para.paragraph_format.line_spacing.pt,sep=":")


    #大纲等级
    if para.style.name == "Normal":
        print("正文")
    elif (para.style.name).startswith("Subtitle"):
        print(para.style.name.replace("Subtitle","副标题"))
    elif (para.style.name).startswith("Heading"):
        print(para.style.name.replace("Heading","标题"))
    else:
        print(para.style.name)


    #段落字单元解析
    print("字单元个数",len(para.runs),sep=":")
    for index,unit in enumerate(para.runs):
        print("字单元"+str(index))

        print("正文", unit.text, sep=":")
        # 字体大小
        if unit.font.size is None:
            print("未设置字体大小")
        else:
            print("字体大小", unit.font.size.pt, sep=":")
        # 判断是否为加粗
        print("加粗", unit.font.bold, sep=":")
        # 是否为斜体
        print("斜体", unit.font.italic, sep=":")
        # 字体名称
        print("字体", unit.font.name, sep=":")
        # 字体颜色
        print("文字颜色", unit.font.color.rgb, sep=":")
        # 是否有下划线
        print("文字下划线", unit.font.underline, sep=":")
        # 是否突出显示
        print("突出显示", unit.font.outline, sep=":")

    print("-----------------------------------")


def wpsParser(filepath):
    print(filepath)

if __name__ == '__main__':
    eachPath("./specificTest")

